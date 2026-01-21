"""
OCR Service for InsightSheet-lite / Meldra. Generic: works with any form, invoice, or document image.
Extracts text from images (JPG, PNG, WebP, BMP, TIFF, GIF) and exports to editable DOC or PDF.
- Form mode: structure (sections, labels, tables, checkboxes) in document flow. Form-agnostic.
- Layout mode: places text at original (x,y) so output matches the image layout. Preserves format.
"""
import io
import os
import re
import logging
from typing import Union, BinaryIO, List, Dict, Any, Tuple, Optional

import httpx

logger = logging.getLogger(__name__)

# OCR.space: 25k req/mo free, 1MB file limit, 2–5s. Set OCR_SPACE_API_KEY to use.
OCR_SPACE_API_URL = "https://api.ocr.space/parse/image"
OCR_SPACE_MAX_BYTES = 1024 * 1024  # 1MB free tier

# Optional imports - fail gracefully if not available
try:
    import pytesseract
    from pytesseract import Output
    from PIL import Image, ImageEnhance
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    Output = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Pt = None
    OxmlElement = None
    qn = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RlTable, TableStyle
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    pdf_canvas = None

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    fitz = None


# Placeholder for fillable "field" so user can type over it in Word/PDF
FIELD_PLACEHOLDER = "________________"

# Single tokens that are labels, not a single checkbox option (generic).
_LABEL_LIKE_SINGLE = frozenset(
    {"id", "name", "email", "phone", "address", "date", "signature", "role", "company", "number", "code", "ref", "tax", "vat"}
)

# Two-word compound labels: never treat as checkbox options when they appear after a label (generic across forms).
_COMPOUND_LABEL_PHRASES = frozenset({
    "employee id", "full name", "form id", "start date", "end date", "phone number", "postal code",
    "street address", "work history", "personal details", "contact number", "email address",
    "company name", "job title", "date of birth", "marital status", "account number", "reference number",
    "invoice number", "order number", "zip code", "tax id", "vat number", "signature date",
    "place of birth", "nationality", "bank account", "sort code", "policy number", "license number",
    "passport number", "national insurance", "social security", "country code", "area code",
})


def _normalize_ocr_text(s: str) -> str:
    """
    Fix common OCR mistakes. Generic: checkboxes [) [|] -> [ ], stray | -> I,
    "Ltd" misread as " d", spaced hyphens (Part - time -> Part-time, ID - 001 -> ID-001), section brackets.
    """
    if not s or not isinstance(s, str):
        return s
    # [| confirm / [| consent -> I confirm / I consent
    s = re.sub(r'\[\s*\|(\s+[a-z])', r'I\1', s)
    # [) [|] [ )] etc (checkbox misreads) -> [ ] (do not consume trailing space)
    s = re.sub(r'\[\s*[\)\|]\]?', '[ ]', s)
    # | at word start when followed by lowercase: "| confirm" -> "I confirm", ". | consent" -> ". I consent"
    s = re.sub(r'(\s|^)\|(\s*[a-z])', r'\1I\2', s)
    # "Acme Logistics d" / "Company d" at end of word -> "Ltd" (common OCR misread)
    s = re.sub(r'(\b\w{2,}) d\b', r'\1 Ltd', s)
    # "Part - time" -> "Part-time", "ONB - 2026 - 001" -> "ONB-2026-001" (OCR.space extra spaces)
    s = re.sub(r'(\w)\s+-\s+(\w)', r'\1-\2', s)
    # "[ 1) Section" -> "1) Section" (stray [ before section header)
    s = re.sub(r'^\[\s*(\d+[\)\.]\s*\S)', r'\1', s)
    return s


def _is_border_or_noise(s: str) -> bool:
    """True if the line is likely a table border, rule, or OCR garbage. Keeps labels and real fields."""
    if not s or not isinstance(s, str):
        return True
    t = s.strip()
    if not t:
        return True
    # Never drop label-like lines: "Full Name:", "Label:", "ID:"
    if re.search(r'\b[a-zA-Z]{2,}\s*:', t):
        return False
    # Mostly dashes, pipes, dots, underscores, equals
    if re.match(r'^[\s\-=\|\.\*#_]+$', t):
        return True
    # Very short non-word only: 1–2 chars with no real word (drop "|", "a", "a ="; keep "ID", "No", "N/A")
    if len(t) <= 2 and not re.search(r'[a-zA-Z]{2,}', t):
        return True
    # Lines that are mostly dashes/equals with no real words (e.g. "a a -----", table borders)
    if re.search(r'[-=]{3,}', t) and not re.search(r'[a-zA-Z]{2,}', t):
        return True
    return False


def _split_into_columns(line: str) -> List[str]:
    """Split a line by tabs or 2+ spaces into columns."""
    if '\t' in line:
        return [c.strip() for c in line.split('\t') if c.strip()]
    return [c.strip() for c in re.split(r'  +', line) if c.strip()]


def _is_section_header(text: str) -> bool:
    """True if line is a section header: "1) Title", "2. Section Name", etc. Generic."""
    if not text or not isinstance(text, str):
        return False
    return bool(re.match(r'^\d+[\)\.]\s*\S', text.strip()))


def _is_label_line(text: str) -> bool:
    """True if line is a form label (ends with :, short, not a section header)."""
    if not text or not isinstance(text, str):
        return False
    t = text.strip()
    if len(t) >= 60:
        return False
    if not t.endswith(':'):
        return False
    if _is_section_header(t):
        return False
    return True


def _is_checkbox_line(text: str) -> Tuple[bool, int]:
    """True if line is checkbox options. Returns (is_checkbox, n_options). Kept for backward compat."""
    ok, opts = _get_checkbox_options(text, False)
    return (ok, len(opts))


def _get_checkbox_options(text: str, prev_line_is_label: bool = False) -> Tuple[bool, List[str], Optional[str]]:
    """
    Detect checkbox lines and return (is_checkbox, options, label). Generic for any form.
    - "[ ] A [ ] B [ ] C" -> (True, [A, B, C], None)
    - "Label: Option A Option B Option C" -> (True, [Option A, Option B, Option C], "Label:")
    - "Option A Option B" when prev_line_is_label -> (True, [Option A, Option B], None)
    """
    if not text or not isinstance(text, str):
        return False, [], None
    t = text.strip()
    if not t or len(t) > 150:
        return False, [], None

    def _reject_single_label(opts: List[str]) -> bool:
        return len(opts) == 1 and (opts[0] or "").lower().strip() in _LABEL_LIKE_SINGLE

    # Explicit [ ] pattern: "A [ ] B [ ] C" or "[ ] A [ ] B"
    if '[ ]' in t:
        parts = [p.strip() for p in re.split(r'\s*\[\s*\]\s*', t) if p.strip()]
        if 2 <= len(parts) <= 6 and all(len(p) < 35 for p in parts) and not _reject_single_label(parts):
            return True, parts, None

    # ☐ (U+2610) or "O O O" style
    if '\u2610' in t or '\u25A1' in t or re.search(r'\b[Oo]\s+[Oo]\s+[Oo]\b', t):
        parts = [p.strip() for p in re.split(r'[\u2610\u25A1\s]+', t) if p.strip() and len(p) > 1]
        if 2 <= len(parts) <= 6 and not _reject_single_label(parts):
            return True, parts, None

    # "Label: Option A Option B Option C" — label and options on one line
    if ':' in t:
        idx = t.index(':')
        left, right = t[: idx + 1].strip(), t[idx + 1:].strip()
        words = [w for w in right.split() if len(w) > 1 and not re.match(r'^[Oo\u2610\u25A1]+$', w)]
        if 2 <= len(words) <= 5 and all(len(w) < 25 for w in words) and not _reject_single_label(words):
            return True, words, left

    # "Option A Option B" when previous line was a label. Generic: reject section-like parentheticals
    # and 2-word compound labels (e.g. "Employee ID", "Work History") via blocklist.
    if prev_line_is_label:
        if '(' in t or ')' in t:
            return False, [], None
        words = [w for w in t.split() if len(w) > 1 and not re.match(r'^[Oo\u2610\u25A1\[\]\)\(]+$', w)]
        if 2 <= len(words) <= 5 and all(len(w) < 25 for w in words) and len(t) < 100 and not _reject_single_label(words):
            if len(words) == 2 and ' '.join(w.lower() for w in words) in _COMPOUND_LABEL_PHRASES:
                return False, [], None
            return True, words, None

    return False, [], None


def _compute_section_boxes(
    layout: List[Dict[str, Any]],
    tables: List[Dict[str, Any]],
    pad_px: int = 6,
) -> List[Dict[str, Any]]:
    """
    Compute one box per section (1), 2), 3)...) around header + all content.
    Generic: works for any form. Returns [{left, top, width, height}, ...].
    """
    if not layout:
        return []
    headers = [ln for ln in layout if _is_section_header((ln.get('text') or '').strip())]
    if not headers:
        return []

    def _key(ln):
        return (ln.get('block_num', 0), ln.get('line_num', 0))

    boxes = []
    for i, h in enumerate(headers):
        start = _key(h)
        end = _key(headers[i + 1]) if i + 1 < len(headers) else (99999, 99999)
        bbox_list = []
        for ln in layout:
            if start <= _key(ln) < end:
                bbox_list.append((ln.get('left', 0), ln.get('top', 0), ln.get('width', 0), ln.get('height', 0)))
        for t in (tables or []):
            k = (t.get('block_num', 0), t.get('line_start', 0))
            if start <= k < end:
                bbox_list.append((t.get('left', 0), t.get('top', 0), t.get('width', 0), t.get('height', 0)))
        if not bbox_list:
            continue
        L = max(0, min(l for l, _, _, _ in bbox_list) - pad_px)
        T = max(0, min(t for _, t, _, _ in bbox_list) - pad_px)
        R = max(l + w for l, _, w, _ in bbox_list) + pad_px
        B = max(t + h for _, t, _, h in bbox_list) + pad_px
        boxes.append({'left': L, 'top': T, 'width': max(1, R - L), 'height': max(1, B - T)})
    return boxes


def _detect_tables_from_words(
    words: List[Dict[str, Any]],
    gap_px: int = 28,
) -> List[Dict[str, Any]]:
    """
    Cluster words into rows by line_num, split into columns by horizontal gaps.
    Returns: [{ "rows": [[c1,c2,...],...], "block_num", "line_start", "line_end", "left", "top", "width", "height" }]
    """
    if not words:
        return []

    # Group by (block_num, line_num), sort
    by_line: Dict[Tuple[int, int], List[Dict]] = {}
    for w in words:
        key = (w.get('block_num', 0), w.get('line_num', 0))
        by_line.setdefault(key, []).append(w)

    tables: List[Dict[str, Any]] = []
    sorted_keys = sorted(by_line.keys())
    i = 0
    while i < len(sorted_keys):
        key = sorted_keys[i]
        line_words = sorted(by_line[key], key=lambda w: w.get('left', 0))
        block_num, line_num = key

        # Column boundaries: gap > gap_px
        cells: List[str] = []
        curr: List[str] = []
        prev_right = -1000
        for w in line_words:
            left = w.get('left', 0)
            width = w.get('width', 0)
            txt = (w.get('text') or '').strip()
            if left - prev_right > gap_px and curr:
                cells.append(_normalize_ocr_text(' '.join(curr)))
                curr = []
            curr.append(txt)
            prev_right = left + width
        if curr:
            cells.append(_normalize_ocr_text(' '.join(curr)))

        # Need 2+ columns to be a table row
        if len(cells) < 2:
            i += 1
            continue

        table_rows = [cells]
        table_keys = [key]
        j = i + 1
        while j < len(sorted_keys):
            nkey = sorted_keys[j]
            nblock, nline = nkey
            nwords = sorted(by_line[nkey], key=lambda w: w.get('left', 0))
            ncells = []
            ncurr = []
            npr = -1000
            for w in nwords:
                left = w.get('left', 0)
                width = w.get('width', 0)
                txt = (w.get('text') or '').strip()
                if left - npr > gap_px and ncurr:
                    ncells.append(_normalize_ocr_text(' '.join(ncurr)))
                    ncurr = []
                ncurr.append(txt)
                npr = left + width
            if ncurr:
                ncells.append(_normalize_ocr_text(' '.join(ncurr)))

            if len(ncells) < 2:
                break
            if abs(len(ncells) - len(cells)) > 1:
                break
            table_rows.append(ncells)
            table_keys.append(nkey)
            j += 1

        if len(table_rows) >= 2:
            # Normalize row lengths
            max_cols = max(len(r) for r in table_rows)
            for r in table_rows:
                while len(r) < max_cols:
                    r.append("")
            # Replace OCR garbage in header row (e.g. "a", "a =" from misread table borders)
            for j, c in enumerate(table_rows[0]):
                cs = (c or "").strip()
                if len(cs) <= 3 and re.match(r'^[a-zA-Z]\s*=?\s*$', cs):
                    table_rows[0][j] = "—"
            # Bbox from first/last words
            all_ws = []
            for k in table_keys:
                all_ws.extend(by_line[k])
            left = min(w.get('left', 0) for w in all_ws)
            top = min(w.get('top', 0) for w in all_ws)
            right = max(w.get('left', 0) + w.get('width', 0) for w in all_ws)
            bottom = max(w.get('top', 0) + w.get('height', 0) for w in all_ws)
            tables.append({
                "rows": table_rows,
                "block_num": block_num,
                "line_start": line_num,
                "line_end": table_keys[-1][1],
                "left": left,
                "top": top,
                "width": max(1, right - left),
                "height": max(1, bottom - top),
            })
            i = j
        else:
            i += 1

    tables.sort(key=lambda t: (t.get('top', 0), t.get('left', 0)))
    return tables


def _parse_form_like_text(text: str) -> List[Dict[str, Any]]:
    """
    Parse OCR text into structure: section, label (with fill line), table, checkbox, paragraph.
    Generic: form-agnostic. Normalizes OCR noise and drops border/garbage. Uses blocklists to avoid
    misclassifying compound labels or section subheaders as checkbox options.
    """
    blocks: List[Dict[str, Any]] = []
    raw = [ln.rstrip() for ln in text.splitlines()]
    lines = []
    for ln in raw:
        n = _normalize_ocr_text(ln)
        if not _is_border_or_noise(n):
            lines.append(n)

    i = 0
    while i < len(lines):
        ln = lines[i]
        if not ln:
            i += 1
            continue

        # --- Section header: "1) Section A", "2. Section B", etc.
        if re.match(r'^\d+[\).]\s*\S', ln):
            blocks.append({"type": "section", "text": ln})
            i += 1
            continue

        # --- "Label: O O O" on one line -> emit label and checkbox
        if ':' in ln and re.search(r'[Oo]\s+[Oo]\s+[Oo]|[\u2610\u25A1]\s+[\u2610\u25A1]\s+[\u2610\u25A1]', ln):
            idx = ln.index(':')
            left, right = ln[: idx + 1], ln[idx + 1 :].strip()
            if len(left) < 80 and re.search(r'[Oo]\s+[Oo]\s+[Oo]', right):
                blocks.append({"type": "label", "text": left})
                i += 1
                continue

        # --- Checkbox: "[ ] A [ ] B [ ] C" or "O O O", "☐ ☐ ☐"; or "Option A Option B" followed by "O O O"
        parts = [x.strip() for x in re.split(r'\[\s*\]', ln) if x.strip()]
        if 2 <= len(parts) <= 6 and all(len(p) < 30 for p in parts):
            blocks.append({"type": "checkbox", "options": parts})
            i += 1
            continue
        checkbox_match = re.search(r'[Oo]\s+[Oo]\s+[Oo]|[\u2610\u25A1]\s+[\u2610\u25A1]\s+[\u2610\u25A1]', ln)
        if checkbox_match or (re.match(r'^[\sOo\u2610\u25A1]+$', ln) and len(ln.strip()) >= 3):
            # Try to get option labels from previous line
            opts: List[str] = []
            if i > 0 and lines[i - 1]:
                prev = lines[i - 1]
                # "Label: O O O" -> label stays, we use generic options
                if prev.endswith(':'):
                    opts = ["Option 1", "Option 2", "Option 3"]
                else:
                    # "Option A Option B" -> use these
                    opts = [w for w in re.split(r'\s+', prev) if len(w) > 1 and not re.match(r'^[Oo\u2610\u25A1]+$', w)]
            if not opts or len(opts) > 6:
                opts = ["Option 1", "Option 2", "Option 3"]
            blocks.append({"type": "checkbox", "options": opts[:5]})
            i += 1
            continue

        # If this line looks like option labels and next is "O O O", we'll handle on next iter.
        # Reject section-like lines with parentheticals (e.g. "Section (Note)") — generic.
        next_is_oo = i + 1 < len(lines) and bool(re.search(r'[Oo]\s+[Oo]\s+[Oo]', lines[i + 1]))
        if next_is_oo and 2 <= len(re.split(r'\s+', ln)) <= 5 and not ln.endswith(':') and '(' not in ln and ')' not in ln:
            opts = [w for w in re.split(r'\s+', ln) if len(w) > 1]
            if opts:
                blocks.append({"type": "checkbox", "options": opts})
                i += 2  # skip next "O O O" line
                continue

        # --- Checkbox options on this line when previous line was a label. Generic: allow 2–5 options;
        # reject section-like parentheticals and 2-word compound labels via blocklist.
        if i > 0 and (lines[i - 1] or '').strip().endswith(':') and not ln.endswith(':') and '(' not in ln and ')' not in ln:
            words = [w for w in ln.split() if len(w) > 1 and not re.match(r'^[Oo\u2610\u25A1\[\]\)\(]+$', w)]
            if 2 <= len(words) <= 5 and all(len(w) < 25 for w in words) and len(ln) < 100:
                if not (len(words) == 2 and ' '.join(w.lower() for w in words) in _COMPOUND_LABEL_PHRASES):
                    blocks.append({"type": "checkbox", "options": words})
                    i += 1
                    continue

        # --- Table: 2+ consecutive lines with same number of columns (2–8) when split by tab or 2+ spaces
        cols = _split_into_columns(ln)
        if len(cols) >= 2 and len(cols) <= 8:
            table_rows = [cols]
            j = i + 1
            while j < len(lines) and lines[j]:
                next_cols = _split_into_columns(lines[j])
                if 2 <= len(next_cols) <= 8 and (len(next_cols) == len(cols) or abs(len(next_cols) - len(cols)) <= 1):
                    table_rows.append(next_cols)
                    j += 1
                else:
                    break
            # Require at least 2 rows to treat as table
            if len(table_rows) >= 2:
                # Normalize row lengths to max
                max_cols = max(len(r) for r in table_rows)
                for r in table_rows:
                    while len(r) < max_cols:
                        r.append("")
                blocks.append({"type": "table", "rows": table_rows})
                i = j
                continue

        # --- Label (ends with :) + fill line. Add an editable underline. Generic.
        if ln.endswith(':') and len(ln) < 100 and not re.match(r'^\d+[\).]', ln):
            blocks.append({"type": "label", "text": ln})
            i += 1
            continue

        # --- Empty or underline-only line after a label: already handled by label. Skip.
        if not ln.strip() or re.match(r'^[\s_\-\.]+$', ln):
            i += 1
            continue

        # --- Normal paragraph
        blocks.append({"type": "paragraph", "text": ln})
        i += 1

    return blocks


def _mime_from_filename(filename: str) -> str:
    ext = (os.path.splitext(filename or "")[1] or "").lower()
    m = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp",
         ".bmp": "image/bmp", ".tiff": "image/tiff", ".tif": "image/tiff", ".gif": "image/gif"}
    return m.get(ext, "image/png")


async def extract_with_layout_ocrspace(
    api_key: str,
    file_content: bytes,
    image_width: int,
    image_height: int,
    filename: str,
) -> Dict[str, Any]:
    """
    Call OCR.space API and return { text, layout, image_width, image_height, tables }
    in the same format as OCRService.extract_with_layout. Uses TextOverlay for coordinates
    when available; otherwise builds a simple layout from ParsedText.
    """
    mime = _mime_from_filename(filename)
    fname = (filename or "image.png").strip() or "image.png"
    data = {"apikey": api_key, "isOverlayRequired": "true", "OCREngine": "2", "language": "eng"}
    files = {"file": (fname, file_content, mime)}

    try:
        async with httpx.AsyncClient() as client:
            r = await client.post(OCR_SPACE_API_URL, data=data, files=files, timeout=30.0)
        r.raise_for_status()
    except httpx.HTTPStatusError as e:
        body = (e.response.text or "")[:400]
        raise RuntimeError(f"OCR.space HTTP {e.response.status_code}: {body}") from e
    except httpx.RequestError as e:
        raise RuntimeError(f"OCR.space request failed: {e}") from e

    j = r.json()

    if j.get("OCRExitCode") != 1:
        raise RuntimeError(f"OCR.space error: {j.get('ErrorMessage', j)}")
    parsed = (j.get("ParsedResults") or [])
    if not parsed:
        raise RuntimeError("OCR.space: no ParsedResults")
    pr = parsed[0]
    parsed_text = (pr.get("ParsedText") or "").strip()
    overlay = pr.get("TextOverlay") or {}
    lines_raw = overlay.get("Lines") or []

    words: List[Dict[str, Any]] = []
    layout: List[Dict[str, Any]] = []

    if lines_raw:
        for line_idx, line in enumerate(lines_raw):
            wrds = line.get("Words") or []
            if not wrds:
                continue
            texts = [str(w.get("WordText") or "").strip() for w in wrds if (w.get("WordText") or "").strip()]
            line_text = _normalize_ocr_text(" ".join(texts))
            if _is_border_or_noise(line_text):
                continue
            lefts = [int(w.get("Left", 0)) for w in wrds]
            tops = [int(w.get("Top", 0)) for w in wrds]
            rights = [int(w.get("Left", 0)) + int(w.get("Width", 0)) for w in wrds]
            bots = [int(w.get("Top", 0)) + int(w.get("Height", 0)) for w in wrds]
            L, T = min(lefts), min(tops)
            R, B = max(rights), max(bots)
            layout.append({
                "text": line_text,
                "left": L,
                "top": T,
                "width": max(1, R - L),
                "height": max(1, B - T),
                "line_num": line_idx,
                "block_num": 0,
            })
            for w in wrds:
                wt = (w.get("WordText") or "").strip()
                if wt:
                    words.append({
                        "text": wt,
                        "left": int(w.get("Left", 0)),
                        "top": int(w.get("Top", 0)),
                        "width": int(w.get("Width", 0)),
                        "height": int(w.get("Height", 0)),
                        "line_num": line_idx,
                        "block_num": 0,
                    })
    else:
        for i, ln in enumerate(parsed_text.splitlines()):
            t = _normalize_ocr_text(ln.strip())
            if not t or _is_border_or_noise(t):
                continue
            layout.append({
                "text": t,
                "left": 0,
                "top": i * 20,
                "width": max(100, min(400, len(t) * 8)),
                "height": 18,
                "line_num": i,
                "block_num": 0,
            })

    layout.sort(key=lambda x: (x.get("top", 0), x.get("left", 0)))
    tables = _detect_tables_from_words(words) if words else []
    text = "\n".join(ln["text"] for ln in layout) if layout else _normalize_ocr_text(parsed_text)

    return {
        "text": text,
        "layout": layout,
        "image_width": image_width,
        "image_height": image_height,
        "tables": tables,
    }


def pdf_from_image(image_bytes: bytes) -> bytes:
    """
    Create a PDF that looks exactly like the input image (full-page image).
    Use for preserve_image / "Exact copy" export. Requires PyMuPDF.
    """
    if not FITZ_AVAILABLE or fitz is None:
        raise RuntimeError("PyMuPDF is not installed. Install: pip install PyMuPDF")
    doc = fitz.open()
    page = doc.new_page()  # A4
    page.insert_image(page.rect, stream=image_bytes)
    buf = io.BytesIO()
    doc.save(buf, deflate=False)
    doc.close()
    return buf.getvalue()


class OCRService:
    """OCR: image -> text, and text -> editable DOC/PDF. Generic for any form, invoice, or document."""

    # Image extensions we support
    ALLOWED_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif', '.gif'}

    def __init__(self):
        if not PYTESSERACT_AVAILABLE:
            raise RuntimeError(
                "OCR dependencies missing. Install: pip install pytesseract Pillow. "
                "Also install Tesseract: https://github.com/tesseract-ocr/tesseract (e.g. apt-get install tesseract-ocr)"
            )

    def extract_text(self, image_data: Union[bytes, BinaryIO]) -> str:
        """
        Run OCR on an image and return extracted text.

        Args:
            image_data: Raw image bytes or file-like object.

        Returns:
            Extracted text as string.
        """
        if not PYTESSERACT_AVAILABLE:
            raise RuntimeError("pytesseract or Pillow not installed.")

        data = image_data.read() if hasattr(image_data, 'read') else image_data
        img = Image.open(io.BytesIO(data))

        # Convert to RGB if necessary (e.g. RGBA, P)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')

        try:
            # PSM 6 = Assume a single uniform block of text (default for documents)
            # PSM 4 = Assume a single column of variable-size text (helps forms with multiple blocks)
            # PSM 3 = Fully automatic; can help noisy/form layouts
            text = pytesseract.image_to_string(img, lang='eng', config='--psm 6')
            return (text or '').strip()
        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract OCR is not installed or not in PATH.")
            raise RuntimeError(
                "Tesseract OCR is not installed. Please install Tesseract: "
                "https://github.com/tesseract-ocr/tesseract - e.g. 'apt-get install tesseract-ocr' on Linux, "
                "or download from GitHub on Windows."
            )
        except Exception as e:
            logger.error(f"OCR extraction error: {e}")
            raise

    def extract_with_layout(self, image_data: Union[bytes, BinaryIO]) -> Dict[str, Any]:
        """
        Run OCR and return text plus per-line bounding boxes so export can preserve layout.
        Returns: { "text": str, "layout": [{"text", "left", "top", "width", "height"}], "image_width": int, "image_height": int }
        """
        if not PYTESSERACT_AVAILABLE:
            raise RuntimeError("pytesseract or Pillow not installed.")

        data = image_data.read() if hasattr(image_data, 'read') else image_data
        img = Image.open(io.BytesIO(data))
        iw, ih = img.size

        # Resize: 1200px to finish within 55s on Railway. 1400px timed out; 1000px dropped fields.
        max_dim = 1200
        if max(iw, ih) > max_dim:
            ratio = max_dim / max(iw, ih)
            new_w, new_h = int(iw * ratio), int(ih * ratio)
            img = img.resize((new_w, new_h), getattr(getattr(Image, 'Resampling', None), 'LANCZOS', None) or Image.LANCZOS)
            iw, ih = img.size

        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        # Preprocess: grayscale and slight contrast boost for better form/table OCR
        if img.mode == 'RGB':
            img = img.convert('L')
        try:
            enh = ImageEnhance.Contrast(img)
            img = enh.enhance(1.25)
        except Exception:
            pass

        try:
            # PSM 4 = single column of variable-sized text (helps forms with sections/tables)
            d = pytesseract.image_to_data(img, lang='eng', config='--psm 4', output_type=Output.DICT)
        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract OCR is not installed or not in PATH.")
            raise RuntimeError(
                "Tesseract OCR is not installed. Please install Tesseract: "
                "https://github.com/tesseract-ocr/tesseract"
            )

        n = len(d['text'])
        words: List[Dict[str, Any]] = []
        lines_map: Dict[Tuple[int, int], Dict[str, Any]] = {}

        for i in range(n):
            t = (d['text'][i] or '').strip()
            if not t:
                continue
            key = (d['block_num'][i], d['line_num'][i])
            words.append({
                'text': t, 'left': d['left'][i], 'top': d['top'][i],
                'width': d['width'][i], 'height': d['height'][i],
                'line_num': d['line_num'][i], 'block_num': d['block_num'][i],
            })
            if key not in lines_map:
                lines_map[key] = {'texts': [], 'left': [], 'top': [], 'width': [], 'height': []}
            lines_map[key]['texts'].append(t)
            lines_map[key]['left'].append(d['left'][i])
            lines_map[key]['top'].append(d['top'][i])
            lines_map[key]['width'].append(d['width'][i])
            lines_map[key]['height'].append(d['height'][i])

        tables = _detect_tables_from_words(words)

        layout = []
        for key in sorted(lines_map.keys()):
            arr = lines_map[key]
            raw_text = ' '.join(arr['texts'])
            text = _normalize_ocr_text(raw_text)
            if _is_border_or_noise(text):
                continue
            left = min(arr['left'])
            top = min(arr['top'])
            right = max(l + w for l, w in zip(arr['left'], arr['width']))
            bottom = max(t + h for t, h in zip(arr['top'], arr['height']))
            layout.append({
                'text': text,
                'left': int(left),
                'top': int(top),
                'width': int(right - left),
                'height': int(bottom - top),
                'line_num': key[1],
                'block_num': key[0],
            })

        # Top-to-bottom, left-to-right order for correct reading and export alignment
        layout.sort(key=lambda ln: (ln.get('top', 0), ln.get('left', 0)))

        # Reconstruct text for backward compatibility (form mode / editing); use normalized
        text = '\n'.join(ln['text'] for ln in layout) if layout else _normalize_ocr_text(
            pytesseract.image_to_string(img, lang='eng', config='--psm 4').strip()
        )

        return {
            'text': text,
            'layout': layout,
            'image_width': iw,
            'image_height': ih,
            'tables': tables,
        }

    def _layout_to_page_scale(self, image_width: int, image_height: int) -> Tuple[float, float, float, float]:
        """Fit image to letter; return (scale, page_w_pt, page_h_pt, _). scale keeps aspect."""
        pw, ph = letter
        if not image_width or not image_height:
            return 1.0, float(pw), float(ph), 1.0
        scale_x = pw / image_width
        scale_y = ph / image_height
        scale = min(scale_x, scale_y)
        return scale, image_width * scale, image_height * scale, scale

    def text_to_pdf_layout(
        self,
        layout: List[Dict[str, Any]],
        image_width: int,
        image_height: int,
        title: str = "OCR Document",
        tables: Optional[List[Dict[str, Any]]] = None,
    ) -> bytes:
        """
        PDF with text placed at original (x,y). Tables are drawn as grids with cell text.
        """
        if not REPORTLAB_AVAILABLE or pdf_canvas is None:
            raise RuntimeError("reportlab is not installed. Install: pip install reportlab")

        scale, page_w, page_h, _ = self._layout_to_page_scale(image_width, image_height)
        buf = io.BytesIO()
        c = pdf_canvas.Canvas(buf, pagesize=(page_w, page_h))
        c.setTitle(title or "OCR Document")
        tables = tables or []
        drawn = set()  # (block_num, line_start)
        c.setStrokeColorRGB(0, 0, 0)

        # Draw section boundary boxes first (each section like 1), 2), 3) as one block — exact match to form)
        for box in _compute_section_boxes(layout, tables):
            x = box['left'] * scale
            y = page_h - (box['top'] + box['height']) * scale
            c.rect(x, y, box['width'] * scale, box['height'] * scale)

        def _find_table(bnum: int, lnum: int):
            for t in tables:
                if t.get('block_num') == bnum and t.get('line_start', 0) <= lnum <= t.get('line_end', 0):
                    return t
            return None

        def _prev_is_label(prev: Optional[Dict]) -> bool:
            return bool(prev and ((prev.get('text') or '').strip().endswith(':')))

        def _draw_table(t: Dict[str, Any]):
            rows = t.get('rows') or []
            if not rows:
                return
            nr, nc = len(rows), max(len(r) for r in rows) if rows else 0
            if nc == 0:
                return
            left = t.get('left', 0) * scale
            top_img = t.get('top', 0)
            w = max(1, t.get('width', 1)) * scale
            h = max(1, t.get('height', 1)) * scale
            row_h = h / nr
            col_w = w / nc
            # PDF y: image top is at page_h - top_img*scale; bottom of table is at page_h - (top_img+h)*scale
            for ri, row in enumerate(rows):
                for cj, cell in enumerate(row):
                    if cj >= nc:
                        break
                    x = left + cj * col_w + 2
                    y_bottom = page_h - (top_img + (ri + 1) * row_h) * scale
                    c.setFont("Helvetica", max(6, min(10, row_h * 0.6)))
                    c.drawString(x, y_bottom, (str(cell) or "")[:80])
            # Grid
            for i in range(nc + 1):
                cx = left + i * col_w
                c.line(cx, page_h - (top_img + h) * scale, cx, page_h - top_img * scale)
            for i in range(nr + 1):
                cy = page_h - (top_img + i * row_h) * scale
                c.line(left, cy, left + w, cy)

        for idx, ln in enumerate(layout):
            bnum = ln.get('block_num', 0)
            lnum = ln.get('line_num', 0)
            t = _find_table(bnum, lnum)
            if t:
                key = (t.get('block_num', 0), t.get('line_start', 0))
                if key not in drawn:
                    _draw_table(t)
                    drawn.add(key)
                continue
            left = ln.get('left', 0)
            top = ln.get('top', 0)
            width = ln.get('width', 0)
            height = ln.get('height', 12)
            text = (ln.get('text') or '').strip()
            if not text:
                continue
            prev_ln = layout[idx - 1] if idx > 0 else None
            x_pt = left * scale
            y_pt = page_h - (top + height) * scale
            font_size = max(8, min(14, height * scale * 0.7))
            c.setFont("Helvetica", font_size)
            is_cb, opts, label = _get_checkbox_options(text, _prev_is_label(prev_ln))

            if _is_section_header(text):
                # Tight rectangular box around section header (same-to-same as uploaded form)
                pad = 2
                rx = (left - pad) * scale
                ry = page_h - (top + height + pad) * scale
                rw = (width + 2 * pad) * scale
                rh = (height + 2 * pad) * scale
                c.rect(rx, ry, rw, rh)
                c.drawString(x_pt, y_pt, text[:500])
            elif is_cb and opts:
                # Label (if any) + [ ] Option for each — preserves alignment. Generic.
                sq, char_pt = 8, 5
                cx = x_pt
                if label:
                    c.drawString(cx, y_pt, (label + " ")[:60])
                    cx += max(50, (len(label) + 1) * char_pt)
                for part in opts:
                    c.rect(cx, y_pt - 6, sq, sq)
                    cx += sq + 4
                    c.drawString(cx, y_pt, (part or "")[:40])
                    cx += max(len(part or "") * char_pt, 40) + 10
            elif _is_label_line(text):
                c.drawString(x_pt, y_pt, text[:500])
                # Generic: extend underline to right margin (exact match to any form's input line)
                ul_len = max(100, min(400, page_w - x_pt - 24))
                c.line(x_pt, y_pt - 2, x_pt + ul_len, y_pt - 2)
            else:
                c.drawString(x_pt, y_pt, text[:500])

        c.save()
        buf.seek(0)
        return buf.read()

    def text_to_docx_layout(
        self,
        layout: List[Dict[str, Any]],
        image_width: int,
        image_height: int,
        title: str = "OCR Document",
        tables: Optional[List[Dict[str, Any]]] = None,
    ) -> bytes:
        """
        DOCX that approximates original positions. Tables are added as real Word tables.
        """
        if not DOCX_AVAILABLE or Pt is None:
            raise RuntimeError("python-docx is not installed. Install: pip install python-docx")

        scale, page_w, page_h, _ = self._layout_to_page_scale(image_width, image_height)
        doc = Document()
        doc.add_heading(title or "OCR Document", level=0)

        tables = tables or []
        drawn = set()

        def _find_table(bnum: int, lnum: int):
            for t in tables:
                if t.get('block_num') == bnum and t.get('line_start', 0) <= lnum <= t.get('line_end', 0):
                    return t
            return None

        prev_bottom = 0
        for idx, ln in enumerate(layout):
            bnum = ln.get('block_num', 0)
            lnum = ln.get('line_num', 0)
            t = _find_table(bnum, lnum)
            if t:
                key = (t.get('block_num', 0), t.get('line_start', 0))
                if key not in drawn:
                    rows = t.get('rows') or []
                    if rows:
                        nc = max(len(r) for r in rows)
                        tbl = doc.add_table(rows=len(rows), cols=nc)
                        tbl.style = 'Table Grid'
                        for ri, row in enumerate(rows):
                            for cj, cell in enumerate(row):
                                if cj < nc:
                                    tbl.rows[ri].cells[cj].text = (str(cell) or "").strip()
                    drawn.add(key)
                continue
            text = (ln.get('text') or '').strip()
            if not text:
                continue
            prev_ln = layout[idx - 1] if idx > 0 else None
            prev_is_label = bool(prev_ln and ((prev_ln.get('text') or '').strip().endswith(':')))
            left = ln.get('left', 0)
            top = ln.get('top', 0)
            width = ln.get('width', 0)
            height = ln.get('height', 12)
            bottom = top + height
            gap = (top - prev_bottom) * scale if prev_bottom else top * scale
            gap_pt = Pt(max(0, min(gap, 72)))
            font_pt = Pt(max(9, min(12, height * scale * 0.5)))
            is_cb, opts, label = _get_checkbox_options(text, prev_is_label)

            if _is_section_header(text):
                # Section header in a tight bordered box (one-cell table), same as uploaded form
                tbl = doc.add_table(rows=1, cols=1)
                tbl.style = 'Table Grid'
                tc = tbl.rows[0].cells[0]
                tc.text = text
                for run in tc.paragraphs[0].runs:
                    run.font.size = font_pt
                    run.font.name = 'Arial'
                try:
                    tbl.columns[0].width = Pt(max(width + 12, 50) * scale)
                except Exception:
                    pass
                tc.paragraphs[0].paragraph_format.space_before = gap_pt
                tc.paragraphs[0].paragraph_format.space_after = Pt(4)
                # Indent table to match original (left) so each section aligns like the image
                try:
                    if OxmlElement and qn:
                        tblPr = tbl._tbl.find(qn('w:tblPr'))
                        if tblPr is None:
                            tblPr = OxmlElement('w:tblPr')
                            tbl._tbl.insert(0, tblPr)
                        ti = OxmlElement('w:tblInd')
                        ti.set(qn('w:type'), 'dxa')
                        ti.set(qn('w:w'), str(int(max(0, left * scale) * 20)))
                        tblPr.append(ti)
                except Exception:
                    pass
            elif is_cb and opts:
                # Checkbox: label (if any) + ☐ Option for each. Generic.
                p = doc.add_paragraph()
                if label:
                    r0 = p.add_run(label + " ")
                    r0.font.size = font_pt
                    r0.font.name = 'Arial'
                for i, o in enumerate(opts or []):
                    if i > 0:
                        p.add_run("  ")
                    p.add_run("\u2610 " + (o or ""))  # ☐ Option
                for run in p.runs:
                    run.font.size = font_pt
                    run.font.name = 'Arial'
                p.paragraph_format.left_indent = Pt(max(0, left * scale))
                p.paragraph_format.space_before = gap_pt
            elif _is_label_line(text):
                p = doc.add_paragraph()
                r1 = p.add_run(text + " ")
                r1.font.size = font_pt
                r1.font.name = 'Arial'
                n = max(24, min(80, int((image_width - left) * scale / 5)))
                r = p.add_run("_" * n)
                r.underline = True
                r.font.size = font_pt
                r.font.name = 'Arial'
                p.paragraph_format.left_indent = Pt(max(0, left * scale))
                p.paragraph_format.space_before = gap_pt
            else:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.size = font_pt
                r.font.name = 'Arial'
                p.paragraph_format.left_indent = Pt(max(0, left * scale))
                p.paragraph_format.space_before = gap_pt
            prev_bottom = bottom

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def text_to_docx(self, text: str, title: str = "OCR Document") -> bytes:
        """
        Create an editable Word document from text, preserving form-like structure:
        sections, label+underline for fields, tables, and checkbox placeholders.
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Install: pip install python-docx")

        doc = Document()
        doc.add_heading(title, level=0)

        blocks = _parse_form_like_text(text)
        if not blocks:
            doc.add_paragraph("(No text extracted)")

        for b in blocks:
            t = b.get("type")
            if t == "section":
                doc.add_heading(b["text"], level=1)
            elif t == "label":
                p = doc.add_paragraph()
                p.add_run(b["text"] + " ")
                r = p.add_run(FIELD_PLACEHOLDER)
                r.underline = True
            elif t == "checkbox":
                opts = b.get("options", ["Option 1", "Option 2", "Option 3"])
                doc.add_paragraph("  ".join("[ ] " + o for o in opts))
            elif t == "table":
                rows = b.get("rows", [])
                if not rows:
                    continue
                nr, nc = len(rows), max(len(r) for r in rows) if rows else 0
                if nc == 0:
                    continue
                table = doc.add_table(rows=nr, cols=nc)
                table.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for cj, cell_text in enumerate(row):
                        if cj < nc:
                            table.rows[ri].cells[cj].text = (cell_text or "").strip()
            elif t == "paragraph":
                doc.add_paragraph(b.get("text", ""))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def text_to_pdf(self, text: str, title: str = "OCR Document") -> bytes:
        """
        Create an editable PDF from text, preserving form-like structure:
        sections, label+underline for fields, tables, and checkbox placeholders.
        """
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("reportlab is not installed. Install: pip install reportlab")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
            title=title,
        )
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))

        blocks = _parse_form_like_text(text)
        if not blocks:
            story.append(Paragraph("(No text extracted)", styles['Normal']))

        for b in blocks:
            t = b.get("type")
            if t == "section":
                story.append(Paragraph(self._escape(b["text"]), styles['Heading1']))
                story.append(Spacer(1, 6))
            elif t == "label":
                s = self._escape(b["text"]) + " <u>" + FIELD_PLACEHOLDER + "</u>"
                story.append(Paragraph(s, styles['Normal']))
                story.append(Spacer(1, 4))
            elif t == "checkbox":
                opts = b.get("options", ["Option 1", "Option 2", "Option 3"])
                s = "  ".join("[ ] " + self._escape(o) for o in opts)
                story.append(Paragraph(s, styles['Normal']))
                story.append(Spacer(1, 4))
            elif t == "table":
                rows = b.get("rows", [])
                if not rows:
                    continue
                nc = max(len(r) for r in rows)
                # Normalize
                for r in rows:
                    while len(r) < nc:
                        r.append("")
                # ReportLab Table needs list of lists of strings; escape for Paragraph later we use plain
                data = [[self._escape(c or "") for c in row] for row in rows]
                tbl = RlTable(data)
                tbl.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 8))
            elif t == "paragraph":
                s = self._escape(b.get("text", ""))
                story.append(Paragraph(s, styles['Normal']))
                story.append(Spacer(1, 6))

        doc.build(story)
        buf.seek(0)
        return buf.read()

    @staticmethod
    def _escape(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @classmethod
    def is_ocr_available(cls) -> bool:
        """Check if OCR (pytesseract + Tesseract) is available."""
        if not PYTESSERACT_AVAILABLE:
            return False
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
